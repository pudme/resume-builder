import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import sys
import platform
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ats_optimizer import ATSOptimizer
import webbrowser
import atexit
import psutil  # For cross-platform process management

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def get_icon_path():
    """Get the appropriate icon path based on the platform"""
    if platform.system() == 'Darwin':  # macOS
        return resource_path(os.path.join("icons", "resume_builder.icns"))
    else:  # Windows
        return resource_path(os.path.join("icons", "resume_builder.ico"))

def get_template_path():
    """Get the path to the templates directory"""
    return resource_path("templates")

def get_output_path():
    """Get the path to the user's Documents folder"""
    return os.path.expanduser("~/Documents")

def check_single_instance():
    """Check if another instance is already running - cross-platform version"""
    lock_file = os.path.join(os.path.expanduser("~"), "ResumeBuilder", "resume_builder.lock")
    lock_dir = os.path.dirname(lock_file)
    
    # Create directory if it doesn't exist
    os.makedirs(lock_dir, exist_ok=True)
    
    # Check if lock file exists
    if os.path.exists(lock_file):
        try:
            # Try to read the PID from the lock file
            with open(lock_file, 'r') as f:
                pid = int(f.read().strip())
            
            # Check if process is still running using psutil
            if psutil.pid_exists(pid):
                return False
            else:
                # Process is not running, remove stale lock file
                os.remove(lock_file)
        except (ValueError, IOError):
            # Lock file is invalid, remove it
            os.remove(lock_file)
    
    # Create lock file with current process ID
    try:
        with open(lock_file, 'w') as f:
            f.write(str(os.getpid()))
        
        # Register cleanup function
        atexit.register(lambda: os.remove(lock_file) if os.path.exists(lock_file) else None)
        return True
    except IOError:
        return False

class ResumeBuilderGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("ATS-Friendly Resume Builder")
        self.root.geometry("800x600")
        
        # Initialize data storage
        self.resume_data = {
            "personal_info": {},
            "professional_summary": "",
            "work_experience": [],
            "education": [],
            "skills": []
        }
        
        # Initialize ATS optimizer
        self.ats_optimizer = ATSOptimizer()
        
        # Set default output directory to Documents folder
        self.output_dir = get_output_path()
        
        # Create main container
        self.main_container = ttk.Frame(root, padding="10")
        self.main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(self.main_container)
        self.notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Create style for headings
        style = ttk.Style()
        style.configure("Heading.TLabel", font=("Helvetica", 10, "bold"))
        
        # Create tabs
        self.create_personal_info_tab()
        self.create_summary_tab()
        self.create_experience_tab()
        self.create_education_tab()
        self.create_skills_tab()
        self.create_optimize_tab()
        self.create_ai_assist_tab()
        
        # Create buttons frame
        self.button_frame = ttk.Frame(self.main_container)
        self.button_frame.grid(row=1, column=0, pady=10)
        
        # Add buttons
        ttk.Button(self.button_frame, text="Save Resume", command=self.save_resume).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.button_frame, text="Clear All", command=self.clear_all).pack(side=tk.LEFT, padx=5)
        
        # Configure grid weights
        self.main_container.columnconfigure(0, weight=1)
        self.main_container.rowconfigure(0, weight=1)

    def create_personal_info_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="Personal Info")
        
        # Create form fields
        fields = [
            ("Full Name", "full_name"),
            ("Email", "email"),
            ("Phone", "phone"),
            ("Location", "location"),
            ("LinkedIn", "linkedin")
        ]
        
        for i, (label, key) in enumerate(fields):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky=tk.W, pady=5)
            entry = ttk.Entry(frame, width=40)
            entry.grid(row=i, column=1, sticky=tk.W, pady=5)
            entry.bind('<KeyRelease>', lambda e, k=key: self.update_personal_info(k, e.widget.get()))
        
        # Add AI assistance section
        ttk.Label(frame, text="\nAI Assistance", style="Heading.TLabel").grid(row=len(fields), column=0, columnspan=2, sticky=tk.W, pady=(20,5))
        
        # Target role input
        ttk.Label(frame, text="Target Role:").grid(row=len(fields)+1, column=0, sticky=tk.W)
        self.target_role = ttk.Entry(frame, width=40)
        self.target_role.grid(row=len(fields)+1, column=1, sticky=tk.W, pady=5)
        
        # AI assistance button
        ttk.Button(frame, text="Get AI Suggestions", command=lambda: self.get_ai_suggestions("personal_info")).grid(row=len(fields)+2, column=0, columnspan=2, pady=5)

    def create_summary_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="Professional Summary")
        
        ttk.Label(frame, text="Write a brief professional summary (2-3 sentences):").pack(anchor=tk.W)
        self.summary_text = scrolledtext.ScrolledText(frame, height=10, width=70)
        self.summary_text.pack(fill=tk.BOTH, expand=True, pady=5)
        self.summary_text.bind('<KeyRelease>', lambda e: self.update_summary(e.widget.get("1.0", tk.END)))
        
        # Add AI assistance section
        ttk.Label(frame, text="\nAI Assistance", style="Heading.TLabel").pack(anchor=tk.W, pady=(20,5))
        
        # Target role input
        ttk.Label(frame, text="Target Role:").pack(anchor=tk.W)
        self.summary_target_role = ttk.Entry(frame, width=70)
        self.summary_target_role.pack(fill=tk.X, pady=5)
        
        # AI assistance button
        ttk.Button(frame, text="Get AI Suggestions", command=lambda: self.get_ai_suggestions("summary")).pack(pady=5)

    def create_experience_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="Work Experience")
        
        # Experience list
        self.experience_frame = ttk.Frame(frame)
        self.experience_frame.pack(fill=tk.BOTH, expand=True)
        
        # Add experience button
        ttk.Button(frame, text="Add Experience", command=self.add_experience).pack(pady=5)
        
        # Add AI assistance section
        ttk.Label(frame, text="\nAI Assistance", style="Heading.TLabel").pack(anchor=tk.W, pady=(20,5))
        
        # Target role input
        ttk.Label(frame, text="Target Role:").pack(anchor=tk.W)
        self.experience_target_role = ttk.Entry(frame, width=70)
        self.experience_target_role.pack(fill=tk.X, pady=5)
        
        # AI assistance button
        ttk.Button(frame, text="Get AI Suggestions", command=lambda: self.get_ai_suggestions("experience")).pack(pady=5)

    def create_education_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="Education")
        
        # Education list
        self.education_frame = ttk.Frame(frame)
        self.education_frame.pack(fill=tk.BOTH, expand=True)
        
        # Add education button
        ttk.Button(frame, text="Add Education", command=self.add_education).pack(pady=5)
        
        # Add AI assistance section
        ttk.Label(frame, text="\nAI Assistance", style="Heading.TLabel").pack(anchor=tk.W, pady=(20,5))
        
        # Target role input
        ttk.Label(frame, text="Target Role:").pack(anchor=tk.W)
        self.education_target_role = ttk.Entry(frame, width=70)
        self.education_target_role.pack(fill=tk.X, pady=5)
        
        # AI assistance button
        ttk.Button(frame, text="Get AI Suggestions", command=lambda: self.get_ai_suggestions("education")).pack(pady=5)

    def create_skills_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="Skills")
        
        ttk.Label(frame, text="Enter your skills (comma-separated):").pack(anchor=tk.W)
        self.skills_text = scrolledtext.ScrolledText(frame, height=10, width=70)
        self.skills_text.pack(fill=tk.BOTH, expand=True, pady=5)
        self.skills_text.bind('<KeyRelease>', lambda e: self.update_skills(e.widget.get("1.0", tk.END)))
        
        # Add AI assistance section
        ttk.Label(frame, text="\nAI Assistance", style="Heading.TLabel").pack(anchor=tk.W, pady=(20,5))
        
        # Target role input
        ttk.Label(frame, text="Target Role:").pack(anchor=tk.W)
        self.skills_target_role = ttk.Entry(frame, width=70)
        self.skills_target_role.pack(fill=tk.X, pady=5)
        
        # AI assistance button
        ttk.Button(frame, text="Get AI Suggestions", command=lambda: self.get_ai_suggestions("skills")).pack(pady=5)

    def create_optimize_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="Optimize")
        
        # Job description input
        ttk.Label(frame, text="Paste the job description you're applying for:").pack(anchor=tk.W)
        self.job_description = scrolledtext.ScrolledText(frame, height=10, width=70)
        self.job_description.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Optimize button
        ttk.Button(frame, text="Get Optimization Suggestions", command=self.get_optimization_suggestions).pack(pady=5)
        
        # Suggestions display
        ttk.Label(frame, text="Optimization Suggestions:").pack(anchor=tk.W)
        self.suggestions_text = scrolledtext.ScrolledText(frame, height=10, width=70)
        self.suggestions_text.pack(fill=tk.BOTH, expand=True, pady=5)

    def create_ai_assist_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="AI Assistant")
        
        # Introduction text
        intro_text = """
        Get AI assistance with your resume using ChatGPT! This feature will help you:
        • Write a professional summary
        • Improve your job descriptions
        • Enhance your skills section
        • Make your resume more ATS-friendly
        
        Note: You'll need a ChatGPT account to use this feature.
        """
        ttk.Label(frame, text=intro_text, wraplength=700).pack(anchor=tk.W, pady=5)
        
        # Section selection
        ttk.Label(frame, text="Select a section to get AI help with:").pack(anchor=tk.W, pady=5)
        self.ai_section = ttk.Combobox(frame, values=[
            "Professional Summary",
            "Work Experience Description",
            "Skills Section",
            "Education Description"
        ], state="readonly")
        self.ai_section.pack(fill=tk.X, pady=5)
        self.ai_section.set("Professional Summary")
        
        # Context input
        ttk.Label(frame, text="Provide some context about your experience:").pack(anchor=tk.W, pady=5)
        self.context_text = scrolledtext.ScrolledText(frame, height=5, width=70)
        self.context_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Job description input (optional)
        ttk.Label(frame, text="(Optional) Paste the job description you're applying for:").pack(anchor=tk.W, pady=5)
        self.ai_job_description = scrolledtext.ScrolledText(frame, height=5, width=70)
        self.ai_job_description.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Generate prompt button
        ttk.Button(frame, text="Generate ChatGPT Prompt", command=self.generate_chatgpt_prompt).pack(pady=5)
        
        # Generated prompt display
        ttk.Label(frame, text="Copy this prompt and paste it into ChatGPT:").pack(anchor=tk.W, pady=5)
        self.prompt_text = scrolledtext.ScrolledText(frame, height=5, width=70)
        self.prompt_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Open ChatGPT button
        ttk.Button(frame, text="Open ChatGPT in Browser", command=self.open_chatgpt).pack(pady=5)

    def update_personal_info(self, key, value):
        self.resume_data["personal_info"][key] = value

    def update_summary(self, value):
        self.resume_data["professional_summary"] = value.strip()

    def update_skills(self, value):
        self.resume_data["skills"] = [skill.strip() for skill in value.split(",") if skill.strip()]

    def add_experience(self):
        frame = ttk.Frame(self.experience_frame)
        frame.pack(fill=tk.X, pady=5)
        
        # Create entry fields
        fields = [
            ("Company", "company"),
            ("Position", "position"),
            ("Start Date", "start_date"),
            ("End Date", "end_date"),
            ("Description", "description")
        ]
        
        entries = {}
        for i, (label, key) in enumerate(fields):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky=tk.W)
            if key == "description":
                entry = scrolledtext.ScrolledText(frame, height=3, width=40)
            else:
                entry = ttk.Entry(frame, width=40)
            entry.grid(row=i, column=1, sticky=tk.W, pady=2)
            entries[key] = entry
        
        # Add save button
        def save_experience():
            experience = {key: entry.get("1.0", tk.END).strip() if isinstance(entry, scrolledtext.ScrolledText) else entry.get()
                         for key, entry in entries.items()}
            self.resume_data["work_experience"].append(experience)
            frame.destroy()
        
        ttk.Button(frame, text="Save Experience", command=save_experience).grid(row=len(fields), column=0, columnspan=2, pady=5)

    def add_education(self):
        frame = ttk.Frame(self.education_frame)
        frame.pack(fill=tk.X, pady=5)
        
        # Create entry fields
        fields = [
            ("Institution", "institution"),
            ("Degree", "degree"),
            ("Field", "field"),
            ("Graduation Date", "graduation_date"),
            ("GPA", "gpa")
        ]
        
        entries = {}
        for i, (label, key) in enumerate(fields):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky=tk.W)
            entry = ttk.Entry(frame, width=40)
            entry.grid(row=i, column=1, sticky=tk.W, pady=2)
            entries[key] = entry
        
        # Add save button
        def save_education():
            education = {key: entry.get() for key, entry in entries.items()}
            self.resume_data["education"].append(education)
            frame.destroy()
        
        ttk.Button(frame, text="Save Education", command=save_education).grid(row=len(fields), column=0, columnspan=2, pady=5)

    def get_optimization_suggestions(self):
        job_description = self.job_description.get("1.0", tk.END).strip()
        if not job_description:
            messagebox.showwarning("Warning", "Please enter a job description first.")
            return
        
        # Generate resume text for analysis
        resume_text = self._generate_resume_text()
        
        # Get suggestions
        suggestions = self.ats_optimizer.get_optimization_suggestions(resume_text, job_description)
        
        # Display suggestions
        self.suggestions_text.delete("1.0", tk.END)
        for suggestion in suggestions:
            self.suggestions_text.insert(tk.END, f"• {suggestion}\n")

    def _generate_resume_text(self) -> str:
        """Generate a text version of the resume for analysis."""
        text_parts = []
        
        # Add personal info
        info = self.resume_data["personal_info"]
        text_parts.append(f"{info.get('full_name', '')}")
        text_parts.append(f"{info.get('email', '')} {info.get('phone', '')} {info.get('location', '')}")
        
        # Add professional summary
        text_parts.append("\nProfessional Summary")
        text_parts.append(self.resume_data["professional_summary"])
        
        # Add work experience
        text_parts.append("\nProfessional Experience")
        for exp in self.resume_data["work_experience"]:
            text_parts.append(f"{exp.get('company', '')} - {exp.get('position', '')}")
            text_parts.append(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            text_parts.append(exp.get('description', ''))
        
        # Add education
        text_parts.append("\nEducation")
        for edu in self.resume_data["education"]:
            text_parts.append(f"{edu.get('institution', '')} - {edu.get('degree', '')} in {edu.get('field', '')}")
            text_parts.append(f"{edu.get('graduation_date', '')}")
        
        # Add skills
        text_parts.append("\nSkills")
        text_parts.append(", ".join(self.resume_data["skills"]))
        
        return "\n".join(text_parts)

    def save_resume(self):
        # Generate the document
        doc = Document()
        
        # Add personal information
        self._add_personal_info(doc)
        
        # Add professional summary
        self._add_professional_summary(doc)
        
        # Add work experience
        self._add_work_experience(doc)
        
        # Add education
        self._add_education(doc)
        
        # Add skills
        self._add_skills(doc)
        
        # Create save dialog
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_filename = f"resume_{timestamp}"
        
        # Ask user for save location and format
        filetypes = [
            ("Word Document", "*.docx"),
            ("PDF Document", "*.pdf"),
            ("All Files", "*.*")
        ]
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=filetypes,
            initialdir=self.output_dir,  # Start in Documents folder
            initialfile=default_filename,
            title="Save Resume As"
        )
        
        if not filename:  # User cancelled the save dialog
            return
            
        try:
            # Save in the selected format
            if filename.lower().endswith('.docx'):
                doc.save(filename)
                messagebox.showinfo("Success", f"Your resume has been saved as:\n{filename}")
            elif filename.lower().endswith('.pdf'):
                # Handle PDF conversion based on platform
                if platform.system() == 'Darwin':  # macOS
                    try:
                        from docx2pdf import convert
                        convert(filename.replace('.pdf', '.docx'), filename)
                        # Remove the temporary docx file
                        os.remove(filename.replace('.pdf', '.docx'))
                        messagebox.showinfo("Success", f"Your resume has been saved as:\n{filename}")
                    except ImportError:
                        messagebox.showerror("Error", "PDF conversion requires the docx2pdf package. Please install it using:\npip install docx2pdf")
                        return
                else:  # Windows
                    try:
                        from docx2pdf import convert
                        convert(filename.replace('.pdf', '.docx'), filename)
                        # Remove the temporary docx file
                        os.remove(filename.replace('.pdf', '.docx'))
                        messagebox.showinfo("Success", f"Your resume has been saved as:\n{filename}")
                    except ImportError:
                        messagebox.showerror("Error", "PDF conversion requires the docx2pdf package. Please install it using:\npip install docx2pdf")
                        return
            else:
                messagebox.showerror("Error", "Unsupported file format. Please save as .docx or .pdf")
                return
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save resume: {str(e)}")

    def clear_all(self):
        if messagebox.askyesno("Confirm", "Are you sure you want to clear all data?"):
            self.resume_data = {
                "personal_info": {},
                "professional_summary": "",
                "work_experience": [],
                "education": [],
                "skills": []
            }
            
            # Clear all input fields
            for widget in self.main_container.winfo_children():
                if isinstance(widget, ttk.Notebook):
                    for tab in widget.winfo_children():
                        for child in tab.winfo_children():
                            if isinstance(child, (ttk.Entry, scrolledtext.ScrolledText)):
                                child.delete(0, tk.END)
                            elif isinstance(child, ttk.Frame):
                                for grandchild in child.winfo_children():
                                    if isinstance(grandchild, (ttk.Entry, scrolledtext.ScrolledText)):
                                        grandchild.delete(0, tk.END)
            
            # Clear AI assistance fields
            if hasattr(self, 'context_text'):
                self.context_text.delete("1.0", tk.END)
            if hasattr(self, 'ai_job_description'):
                self.ai_job_description.delete("1.0", tk.END)
            if hasattr(self, 'prompt_text'):
                self.prompt_text.delete("1.0", tk.END)
            
            messagebox.showinfo("Success", "All data has been cleared.")

    def _add_personal_info(self, doc):
        """Add personal information to the document"""
        # Add name as title
        name = self.resume_data["personal_info"].get("name", "")
        title = doc.add_heading(name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact information
        contact_info = []
        if "email" in self.resume_data["personal_info"]:
            contact_info.append(self.resume_data["personal_info"]["email"])
        if "phone" in self.resume_data["personal_info"]:
            contact_info.append(self.resume_data["personal_info"]["phone"])
        if "location" in self.resume_data["personal_info"]:
            contact_info.append(self.resume_data["personal_info"]["location"])
        
        if contact_info:
            contact = doc.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.add_run(" | ".join(contact_info))
        
        # Add LinkedIn if available
        if "linkedin" in self.resume_data["personal_info"]:
            linkedin = doc.add_paragraph()
            linkedin.alignment = WD_ALIGN_PARAGRAPH.CENTER
            linkedin.add_run(self.resume_data["personal_info"]["linkedin"])
        
        # Add a line break
        doc.add_paragraph()

    def _add_professional_summary(self, doc):
        """Add professional summary to the document."""
        # Add section heading with custom style
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROFESSIONAL SUMMARY")
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.all_caps = True
        heading.space_after = Pt(8)
        
        # Add summary text with proper spacing
        summary_paragraph = doc.add_paragraph()
        summary_run = summary_paragraph.add_run(self.resume_data["professional_summary"])
        summary_run.font.size = Pt(11)
        summary_paragraph.space_after = Pt(16)

    def _add_work_experience(self, doc):
        """Add work experience to the document."""
        # Add section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROFESSIONAL EXPERIENCE")
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.all_caps = True
        heading.space_after = Pt(8)
        
        for experience in self.resume_data["work_experience"]:
            # Add company and position with proper formatting
            p = doc.add_paragraph()
            company_run = p.add_run(experience.get("company", ""))
            company_run.font.bold = True
            company_run.font.size = Pt(12)
            p.add_run(" • ")  # Changed separator to bullet
            position_run = p.add_run(experience.get("position", ""))
            position_run.font.size = Pt(12)
            p.space_after = Pt(2)
            
            # Add dates in italics
            p = doc.add_paragraph()
            date_run = p.add_run(f"{experience.get('start_date', '')} - {experience.get('end_date', '')}")
            date_run.font.italic = True
            date_run.font.size = Pt(10)
            p.space_after = Pt(6)
            
            # Add description with bullet points
            desc = experience.get("description", "")
            if desc:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12)
                desc_run = p.add_run(desc)
                desc_run.font.size = Pt(11)
                p.space_after = Pt(16)

    def _add_education(self, doc):
        """Add education to the document."""
        # Add section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("EDUCATION")
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.all_caps = True
        heading.space_after = Pt(8)
        
        for education in self.resume_data["education"]:
            # Add institution and degree
            p = doc.add_paragraph()
            institution_run = p.add_run(education.get("institution", ""))
            institution_run.font.bold = True
            institution_run.font.size = Pt(12)
            p.add_run(" • ")  # Changed separator to bullet
            degree_run = p.add_run(f"{education.get('degree', '')} in {education.get('field', '')}")
            degree_run.font.size = Pt(12)
            p.space_after = Pt(2)
            
            # Add dates and GPA in italics
            p = doc.add_paragraph()
            date_run = p.add_run(education.get("graduation_date", ""))
            date_run.font.italic = True
            date_run.font.size = Pt(10)
            if education.get("gpa"):
                date_run.add_run(f" • GPA: {education['gpa']}")  # Changed separator to bullet
            p.space_after = Pt(16)

    def _add_skills(self, doc):
        """Add skills to the document."""
        # Add section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("SKILLS")
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.all_caps = True
        heading.space_after = Pt(8)
        
        # Add skills with proper formatting
        skills_text = " • ".join(self.resume_data["skills"])  # Changed separator to bullet
        p = doc.add_paragraph()
        skills_run = p.add_run(skills_text)
        skills_run.font.size = Pt(11)
        p.space_after = Pt(16)

    def generate_chatgpt_prompt(self):
        section = self.ai_section.get()
        context = self.context_text.get("1.0", tk.END).strip()
        job_description = self.ai_job_description.get("1.0", tk.END).strip()
        
        if not context:
            messagebox.showwarning("Warning", "Please provide some context about your experience.")
            return
        
        # Generate appropriate prompt based on section
        job_desc_part = f"\nThis is for the following job:\n{job_description}\n" if job_description else ""
        
        if section == "Professional Summary":
            prompt = f"""Please help me write a professional summary for my resume. Here's my background:

{context}{job_desc_part}
Please write a concise, ATS-friendly professional summary (2-3 sentences) that highlights my key strengths and experience. Make it specific and impactful, but keep it honest and factual based on the information provided."""
        
        elif section == "Work Experience Description":
            prompt = f"""Please help me write a strong job description for my resume. Here's my background:

{context}{job_desc_part}
Please write a clear, ATS-friendly job description that highlights my achievements and responsibilities. Use action verbs and include specific metrics where possible. Keep it honest and factual based on the information provided."""
        
        elif section == "Skills Section":
            prompt = f"""Please help me create a skills section for my resume. Here's my background:

{context}{job_desc_part}
Please suggest relevant skills that I should include, based on my experience and the job requirements. Include both technical and soft skills. Keep it honest and factual based on the information provided."""
        
        else:  # Education Description
            prompt = f"""Please help me write a strong education section for my resume. Here's my background:

{context}{job_desc_part}
Please write a clear, ATS-friendly education description that highlights my academic achievements and relevant coursework. Keep it honest and factual based on the information provided."""
        
        # Display the prompt
        self.prompt_text.delete("1.0", tk.END)
        self.prompt_text.insert("1.0", prompt)

    def open_chatgpt(self):
        webbrowser.open("https://chat.openai.com")

    def get_ai_suggestions(self, section):
        # Get the target role from the appropriate field
        target_role = ""
        if section == "personal_info":
            target_role = self.target_role.get()
        elif section == "summary":
            target_role = self.summary_target_role.get()
        elif section == "experience":
            target_role = self.experience_target_role.get()
        elif section == "education":
            target_role = self.education_target_role.get()
        elif section == "skills":
            target_role = self.skills_target_role.get()
        
        if not target_role:
            messagebox.showwarning("Warning", "Please enter a target role to get AI suggestions.")
            return
        
        # Get the job description if available
        job_description = self.job_description.get("1.0", tk.END).strip() if hasattr(self, 'job_description') else ""
        
        # Generate appropriate prompt based on section
        context = self._get_section_context(section)
        
        # Create the job description part
        job_desc_part = f"\nThis is for the following job:\n{job_description}\n" if job_description else ""
        
        if section == "personal_info":
            prompt = f"""Please help me optimize my personal information for a {target_role} position. Here's my current information:

{context}{job_desc_part}
Please suggest improvements to make my personal information more professional and relevant for this role. Keep all information factual and honest."""
        
        elif section == "summary":
            prompt = f"""Please help me write a professional summary for a {target_role} position. Here's my background:

{context}{job_desc_part}
Please write a concise, ATS-friendly professional summary (2-3 sentences) that highlights my key strengths and experience. Make it specific and impactful, but keep it honest and factual based on the information provided."""
        
        elif section == "experience":
            prompt = f"""Please help me optimize my work experience for a {target_role} position. Here's my experience:

{context}{job_desc_part}
Please suggest improvements to make my work experience descriptions more impactful and relevant for this role. Use action verbs and include specific metrics where possible. Keep all information factual and honest."""
        
        elif section == "education":
            prompt = f"""Please help me optimize my education section for a {target_role} position. Here's my education:

{context}{job_desc_part}
Please suggest improvements to make my education section more relevant for this role. Highlight relevant coursework and achievements. Keep all information factual and honest."""
        
        else:  # skills
            prompt = f"""Please help me optimize my skills section for a {target_role} position. Here's my current skills:

{context}{job_desc_part}
Please suggest relevant skills that I should include, based on my experience and the job requirements. Include both technical and soft skills. Keep it honest and factual based on the information provided."""
        
        # Display the prompt
        self.prompt_text.delete("1.0", tk.END)
        self.prompt_text.insert("1.0", prompt)
        
        # Switch to AI Assistant tab
        self.notebook.select(self.notebook.index("AI Assistant"))

    def _get_section_context(self, section):
        if section == "personal_info":
            info = self.resume_data["personal_info"]
            return f"Name: {info.get('full_name', '')}\nEmail: {info.get('email', '')}\nPhone: {info.get('phone', '')}\nLocation: {info.get('location', '')}\nLinkedIn: {info.get('linkedin', '')}"
        
        elif section == "summary":
            return self.resume_data["professional_summary"]
        
        elif section == "experience":
            context = []
            for exp in self.resume_data["work_experience"]:
                context.append(f"Company: {exp.get('company', '')}")
                context.append(f"Position: {exp.get('position', '')}")
                context.append(f"Duration: {exp.get('start_date', '')} - {exp.get('end_date', '')}")
                context.append(f"Description: {exp.get('description', '')}\n")
            return "\n".join(context)
        
        elif section == "education":
            context = []
            for edu in self.resume_data["education"]:
                context.append(f"Institution: {edu.get('institution', '')}")
                context.append(f"Degree: {edu.get('degree', '')} in {edu.get('field', '')}")
                context.append(f"Graduation Date: {edu.get('graduation_date', '')}")
                context.append(f"GPA: {edu.get('gpa', '')}\n")
            return "\n".join(context)
        
        else:  # skills
            return ", ".join(self.resume_data["skills"])

if __name__ == "__main__":
    try:
        # Check if another instance is running
        if not check_single_instance():
            messagebox.showerror("Error", "Another instance of Resume Builder is already running.")
            sys.exit(1)
        
        # Create and configure the main window
        root = tk.Tk()
        root.title("Resume Builder")
        
        # Set window size and position
        window_width = 800
        window_height = 600
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # Set window icon
        try:
            icon_path = get_icon_path()
            if os.path.exists(icon_path):
                root.iconbitmap(icon_path)
        except Exception as e:
            print(f"Warning: Could not set window icon: {str(e)}")
        
        # Create the application
        app = ResumeBuilderGUI(root)
        
        # Start the main loop
        root.mainloop()
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")
        sys.exit(1) 