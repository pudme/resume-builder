import os
import json
from datetime import datetime
from rich.console import Console
from rich.prompt import Prompt, Confirm
from rich.panel import Panel
from rich.markdown import Markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ats_optimizer import ATSOptimizer

class ResumeBuilder:
    def __init__(self):
        self.console = Console()
        self.resume_data = {}
        self.templates_dir = "templates"
        self.output_dir = "output"
        self.ats_optimizer = ATSOptimizer()
        self._setup_directories()

    def _setup_directories(self):
        """Create necessary directories if they don't exist."""
        for directory in [self.templates_dir, self.output_dir]:
            if not os.path.exists(directory):
                os.makedirs(directory)

    def start(self):
        """Start the resume builder application."""
        self.console.print(Panel.fit(
            "[bold blue]Welcome to the ATS-Friendly Resume Builder![/]\n\n"
            "This tool will help you create or optimize your resume.",
            title="Resume Builder"
        ))
        
        has_resume = Confirm.ask("Do you already have a resume you'd like to optimize?")
        
        if has_resume:
            self.optimize_existing_resume()
        else:
            self.create_new_resume()

    def optimize_existing_resume(self):
        """Handle optimization of existing resume."""
        self.console.print("\n[yellow]Please provide the path to your resume file (PDF or DOCX)[/]")
        file_path = Prompt.ask("File path")
        
        if not os.path.exists(file_path):
            self.console.print("[red]Error: File not found![/]")
            return

        # Load the document
        doc = Document(file_path)
        
        # Extract text for analysis
        resume_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Ask for job description
        self.console.print("\n[yellow]Please paste the job description you're applying for:[/]")
        job_description = Prompt.ask("Job Description")
        
        # Get optimization suggestions
        suggestions = self.ats_optimizer.get_optimization_suggestions(resume_text, job_description)
        
        # Display suggestions
        self.console.print("\n[bold green]Optimization Suggestions:[/]")
        for suggestion in suggestions:
            self.console.print(f"- {suggestion}")
        
        # Ask if user wants to apply optimizations
        if Confirm.ask("\nWould you like to apply these optimizations to your resume?"):
            # Optimize the document
            optimized_doc = self.ats_optimizer.optimize_document(doc)
            
            # Save the optimized version
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"optimized_resume_{timestamp}.docx"
            output_path = os.path.join(self.output_dir, output_filename)
            optimized_doc.save(output_path)
            
            self.console.print(f"\n[green]Your optimized resume has been saved as: {output_path}[/]")
        else:
            self.console.print("\n[yellow]No changes were made to your resume.[/]")

    def create_new_resume(self):
        """Guide user through creating a new resume."""
        sections = [
            "personal_info",
            "professional_summary",
            "work_experience",
            "education",
            "skills"
        ]
        
        for section in sections:
            self.collect_section_info(section)
        
        # Ask if user wants to optimize for a specific job
        if Confirm.ask("\nWould you like to optimize this resume for a specific job?"):
            self.console.print("\n[yellow]Please paste the job description:[/]")
            job_description = Prompt.ask("Job Description")
            
            # Get optimization suggestions
            resume_text = self._generate_resume_text()
            suggestions = self.ats_optimizer.get_optimization_suggestions(resume_text, job_description)
            
            # Display suggestions
            self.console.print("\n[bold green]Optimization Suggestions:[/]")
            for suggestion in suggestions:
                self.console.print(f"- {suggestion}")
            
            # Ask if user wants to apply suggestions
            if Confirm.ask("\nWould you like to apply these suggestions to your resume?"):
                self.apply_optimization_suggestions(suggestions)
        
        self.generate_resume()

    def _generate_resume_text(self) -> str:
        """Generate a text version of the resume for analysis."""
        text_parts = []
        
        # Add personal info
        info = self.resume_data["personal_info"]
        text_parts.append(f"{info['full_name']}")
        text_parts.append(f"{info['email']} {info['phone']} {info['location']}")
        
        # Add professional summary
        text_parts.append("\nProfessional Summary")
        text_parts.append(self.resume_data["professional_summary"])
        
        # Add work experience
        text_parts.append("\nProfessional Experience")
        for exp in self.resume_data["work_experience"]:
            text_parts.append(f"{exp['company']} - {exp['position']}")
            text_parts.append(f"{exp['start_date']} - {exp['end_date']}")
            text_parts.append(exp["description"])
        
        # Add education
        text_parts.append("\nEducation")
        for edu in self.resume_data["education"]:
            text_parts.append(f"{edu['institution']} - {edu['degree']} in {edu['field']}")
            text_parts.append(f"{edu['graduation_date']}")
        
        # Add skills
        text_parts.append("\nSkills")
        text_parts.append(", ".join(self.resume_data["skills"]))
        
        return "\n".join(text_parts)

    def apply_optimization_suggestions(self, suggestions: list):
        """Apply optimization suggestions to the resume data."""
        for suggestion in suggestions:
            if "required skills" in suggestion.lower():
                # Extract skills from suggestion
                skills_text = suggestion.split(":")[1].strip()
                new_skills = [skill.strip() for skill in skills_text.split(",")]
                
                # Ask user about each missing required skill
                self.console.print("\n[yellow]For each missing required skill, please indicate if you have it:[/]")
                for skill in new_skills:
                    if Confirm.ask(f"Do you have experience with {skill}?"):
                        self.resume_data["skills"].append(skill)
                    else:
                        self.console.print(f"[yellow]Note: {skill} is listed as required. Consider gaining experience in this area.[/]")
            
            elif "preferred skills" in suggestion.lower():
                # Extract skills from suggestion
                skills_text = suggestion.split(":")[1].strip()
                new_skills = [skill.strip() for skill in skills_text.split(",")]
                
                # Ask user about each missing preferred skill
                self.console.print("\n[yellow]For each preferred skill, please indicate if you have it:[/]")
                for skill in new_skills:
                    if Confirm.ask(f"Do you have experience with {skill}?"):
                        self.resume_data["skills"].append(skill)
                    else:
                        self.console.print(f"[yellow]Note: {skill} is listed as preferred but not required.[/]")
            
            elif "experience level" in suggestion.lower():
                # Extract experience level from suggestion
                level = suggestion.split(":")[1].strip()
                
                # Ask user if they want to highlight their experience level
                if Confirm.ask(f"Would you like to highlight your {level} experience level in your summary?"):
                    # Only add if it's not already there
                    if level.lower() not in self.resume_data["professional_summary"].lower():
                        self.resume_data["professional_summary"] = f"{self.resume_data['professional_summary']} With {level} of experience."
            
            elif "education requirements" in suggestion.lower():
                # Extract education requirements from suggestion
                requirements = suggestion.split(":")[1].strip()
                
                # Ask user if they meet these requirements
                if not Confirm.ask(f"Do you meet these education requirements: {requirements}?"):
                    self.console.print("[yellow]Note: You may want to consider how to address this in your application.[/]")

    def collect_section_info(self, section):
        """Collect information for each resume section."""
        self.console.print(f"\n[bold green]Let's work on your {section.replace('_', ' ').title()}[/]")
        
        if section == "personal_info":
            self.resume_data[section] = {
                "full_name": Prompt.ask("Full Name"),
                "email": Prompt.ask("Email"),
                "phone": Prompt.ask("Phone Number"),
                "location": Prompt.ask("Location (City, State)"),
                "linkedin": Prompt.ask("LinkedIn Profile (optional)", default="")
            }
        
        elif section == "professional_summary":
            self.console.print("\n[italic]Write a brief professional summary (2-3 sentences)[/]")
            self.resume_data[section] = Prompt.ask("Summary")
        
        elif section == "work_experience":
            self.resume_data[section] = []
            while Confirm.ask("Add work experience?"):
                experience = {
                    "company": Prompt.ask("Company Name"),
                    "position": Prompt.ask("Position Title"),
                    "start_date": Prompt.ask("Start Date (MM/YYYY)"),
                    "end_date": Prompt.ask("End Date (MM/YYYY) or 'Present'"),
                    "description": Prompt.ask("Job Description")
                }
                self.resume_data[section].append(experience)
        
        elif section == "education":
            self.resume_data[section] = []
            while Confirm.ask("Add education?"):
                education = {
                    "institution": Prompt.ask("Institution Name"),
                    "degree": Prompt.ask("Degree"),
                    "field": Prompt.ask("Field of Study"),
                    "graduation_date": Prompt.ask("Graduation Date (MM/YYYY)"),
                    "gpa": Prompt.ask("GPA (optional)", default="")
                }
                self.resume_data[section].append(education)
        
        elif section == "skills":
            self.console.print("\n[italic]Enter your skills (comma-separated)[/]")
            skills_input = Prompt.ask("Skills")
            self.resume_data[section] = [skill.strip() for skill in skills_input.split(",")]

    def generate_resume(self):
        """Generate the final resume document."""
        doc = Document()
        
        # Add personal information
        self._add_personal_info(doc)
        
        # Add professional summary
        self._add_professional_summary(doc)
        
        # Add work experience
        self._add_work_experience(doc)
        
        # Add education
        self._add_education(doc)
        
        # Add skills
        self._add_skills(doc)
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        
        self.console.print(f"\n[green]Your resume has been generated successfully![/]")
        self.console.print(f"Saved as: {output_path}")

    def _add_personal_info(self, doc):
        """Add personal information to the document."""
        info = self.resume_data["personal_info"]
        
        # Add name
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(info["full_name"])
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact information
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(f"{info['email']} | {info['phone']} | {info['location']}")
        if info["linkedin"]:
            contact_paragraph.add_run(f" | {info['linkedin']}")
        
        doc.add_paragraph()  # Add spacing

    def _add_professional_summary(self, doc):
        """Add professional summary to the document."""
        doc.add_paragraph("Professional Summary", style="Heading 1")
        doc.add_paragraph(self.resume_data["professional_summary"])
        doc.add_paragraph()  # Add spacing

    def _add_work_experience(self, doc):
        """Add work experience to the document."""
        doc.add_paragraph("Professional Experience", style="Heading 1")
        
        for experience in self.resume_data["work_experience"]:
            # Add company and position
            p = doc.add_paragraph()
            p.add_run(experience["company"]).bold = True
            p.add_run(" | ")
            p.add_run(experience["position"])
            
            # Add dates
            p = doc.add_paragraph()
            p.add_run(f"{experience['start_date']} - {experience['end_date']}")
            
            # Add description
            doc.add_paragraph(experience["description"])
            doc.add_paragraph()  # Add spacing

    def _add_education(self, doc):
        """Add education to the document."""
        doc.add_paragraph("Education", style="Heading 1")
        
        for education in self.resume_data["education"]:
            # Add institution and degree
            p = doc.add_paragraph()
            p.add_run(education["institution"]).bold = True
            p.add_run(" | ")
            p.add_run(f"{education['degree']} in {education['field']}")
            
            # Add dates and GPA
            p = doc.add_paragraph()
            p.add_run(f"{education['graduation_date']}")
            if education["gpa"]:
                p.add_run(f" | GPA: {education['gpa']}")
            
            doc.add_paragraph()  # Add spacing

    def _add_skills(self, doc):
        """Add skills to the document."""
        doc.add_paragraph("Skills", style="Heading 1")
        skills_text = ", ".join(self.resume_data["skills"])
        doc.add_paragraph(skills_text)

if __name__ == "__main__":
    builder = ResumeBuilder()
    builder.start() 